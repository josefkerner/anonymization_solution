from pydantic import BaseModel
from typing import List
from docx import Document

class Sentence(BaseModel):
    sentence: str
    start_position: int
    end_position: int

class Paragraph(BaseModel):
    start_position: int
    end_position: int
    text: str

class WordWriter:
    def __init__(self):
        pass

    def write_sentences(self, sentences: List[Sentence], file_path: str):
        document = Document(file_path)
        paragraphs = document.paragraphs

        first_char = 0
        par_objects = []
        for paragraph in paragraphs:
            par: Paragraph = Paragraph(start_position=first_char,
                                       end_position=first_char + len(paragraph.text),
                                       text=paragraph.text)
            par_objects.append(par)
        for sentence in sentences:
            for par in par_objects:
                if sentence.start_position >= par.start_position and sentence.end_position <= par.end_position:
                    paragraph = paragraphs[par_objects.index(par)]
                    run = paragraph.add_run()
                    run.text = sentence.sentence
                    break

